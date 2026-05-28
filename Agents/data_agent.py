import sqlite3
import pandas as pd
from docx import Document
import pdfplumber

from Agents.dataset_profiler import DatasetProfiler


class DataAgent:
    """
    Loads, validates, profiles, and standardizes business datasets.

    Supported file types:
    - CSV
    - Excel
    - SQLite database
    - Word documents with tables
    - PDF documents with structured tables or simple structured rows

    Standard internal schema:
    - date
    - product
    - revenue
    """

    def load_data(self, file_path):
        file_path_lower = file_path.lower()

        if file_path_lower.endswith(".csv"):
            df = pd.read_csv(file_path)

        elif file_path_lower.endswith(".xlsx"):
            df = pd.read_excel(file_path)

        elif file_path_lower.endswith(".db") or file_path_lower.endswith(".sqlite"):
            df = self.load_sqlite(file_path)

        elif file_path_lower.endswith(".docx"):
            df = self.load_docx(file_path)

        elif file_path_lower.endswith(".pdf"):
            df = self.load_pdf(file_path)

        else:
            raise ValueError(
                "Unsupported file format. Please upload CSV, Excel, SQLite, Word, or PDF."
            )

        df = self.validate(df)

        profiler = DatasetProfiler()
        profile = profiler.profile(df)

        print("\n=== DATASET PROFILE ===")
        print(profile)

        df = self.standardize_columns(df)

        return df

    def load_sqlite(self, file_path):
        conn = sqlite3.connect(file_path)

        tables = pd.read_sql_query(
            "SELECT name FROM sqlite_master WHERE type='table';",
            conn
        )

        if tables.empty:
            conn.close()
            raise ValueError("No tables found in SQLite database.")

        best_df = None
        best_table = None

        revenue_keywords = [
            "revenue", "sales", "amount", "total",
            "invoice", "price", "value"
        ]

        entity_keywords = [
            "product", "item", "category", "store",
            "customer", "employee", "artist", "album",
            "track", "country", "city"
        ]

        date_keywords = [
            "date", "time", "created", "invoice", "order"
        ]

        for table_name in tables["name"]:
            df = pd.read_sql_query(f"SELECT * FROM {table_name}", conn)

            columns_lower = [
                str(col).lower().replace(" ", "_")
                for col in df.columns
            ]

            has_revenue = any(
                any(keyword in col for keyword in revenue_keywords)
                for col in columns_lower
            )

            has_entity = any(
                any(keyword in col for keyword in entity_keywords)
                for col in columns_lower
            )

            has_date = any(
                any(keyword in col for keyword in date_keywords)
                for col in columns_lower
            )

            if has_revenue and has_date:
                best_df = df
                best_table = table_name
                break

            if best_df is None and (has_revenue or has_entity or has_date):
                best_df = df
                best_table = table_name

        conn.close()

        if best_df is None:
            raise ValueError("No suitable business table found in SQLite database.")

        print(f"\nSQLite table selected: {best_table}")

        return best_df

    def load_docx(self, file_path):
        document = Document(file_path)

        if not document.tables:
            raise ValueError("No tables found in Word document.")

        table = document.tables[0]
        rows = []

        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])

        if len(rows) < 2:
            raise ValueError("Word table must contain a header row and data rows.")

        headers = rows[0]
        data = rows[1:]

        df = pd.DataFrame(data, columns=headers)

        return df

    def load_pdf(self, file_path):
        all_rows = []

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()

                for table in tables:
                    for row in table:
                        cleaned_row = [
                            str(cell).strip() if cell is not None else ""
                            for cell in row
                        ]

                        if cleaned_row and any(cell for cell in cleaned_row):
                            all_rows.append(cleaned_row)

                if not all_rows:
                    text = page.extract_text()

                    if text:
                        lines = text.split("\n")

                        for line in lines:
                            parts = line.split()

                            if len(parts) >= 3:
                                date_value = parts[0]
                                revenue_value = parts[-1]
                                entity_value = " ".join(parts[1:-1])

                                all_rows.append([
                                    date_value,
                                    entity_value,
                                    revenue_value
                                ])

        if not all_rows:
            raise ValueError(
                "No usable table or structured rows found in PDF document."
            )

        first_row = all_rows[0]
        first_cell = str(first_row[0]) if first_row and first_row[0] is not None else ""

        if (
            "date" in first_cell.lower()
            or "order" in first_cell.lower()
            or "invoice" in first_cell.lower()
        ):
            headers = [str(col) for col in first_row]
            data = all_rows[1:]
        else:
            headers = ["order_date", "store", "weekly_sales"]
            data = all_rows

        max_columns = len(headers)

        cleaned_data = []
        for row in data:
            row = [str(cell).strip() if cell is not None else "" for cell in row]

            if len(row) < max_columns:
                row = row + [""] * (max_columns - len(row))

            if len(row) > max_columns:
                row = row[:max_columns]

            cleaned_data.append(row)

        df = pd.DataFrame(cleaned_data, columns=headers)

        return df

    def validate(self, df):
        if df.empty:
            raise ValueError("Dataset is empty.")

        df = df.dropna(how="all")

        return df

    def normalize_column_name(self, column):
        return (
            str(column)
            .lower()
            .strip()
            .replace(" ", "_")
            .replace("-", "_")
            .replace(".", "_")
        )

    def find_column(self, df, aliases, preferred_type=None):
        normalized_columns = {
            self.normalize_column_name(col): col
            for col in df.columns
        }

        for alias in aliases:
            if alias in normalized_columns:
                return normalized_columns[alias]

        for normalized, original in normalized_columns.items():
            for alias in aliases:
                if alias in normalized:
                    return original

        if preferred_type == "numeric":
            numeric_columns = df.select_dtypes(include=["number"]).columns.tolist()
            if numeric_columns:
                return numeric_columns[0]

            for col in df.columns:
                converted = pd.to_numeric(df[col], errors="coerce")
                if converted.notna().sum() >= max(1, len(df) * 0.6):
                    return col

        if preferred_type == "date":
            for col in df.columns:
                converted = pd.to_datetime(df[col], errors="coerce")
                if converted.notna().sum() >= max(1, len(df) * 0.6):
                    return col

        return None

    def standardize_columns(self, df):
        date_aliases = [
            "date", "order_date", "invoice_date", "transaction_date",
            "sale_date", "sales_date", "created_at", "timestamp",
            "purchase_date", "billing_date", "week", "month"
        ]

        product_aliases = [
            "product", "product_name", "product_line", "item", "item_name",
            "sku", "category", "description", "stock_code", "brand",
            "department", "segment", "store", "branch", "region", "city",
            "customer_type", "sales_channel", "customerid", "customer_id",
            "employeeid", "employee_id", "artistid", "artist_id",
            "albumid", "album_id", "trackid", "track_id", "title"
        ]

        revenue_aliases = [
            "revenue", "sales", "sales_amount", "amount", "total",
            "total_sales", "income", "turnover", "price", "unit_price",
            "weekly_sales", "net_sales", "gross_sales", "value",
            "invoice_amount", "order_value", "profit"
        ]

        date_col = self.find_column(df, date_aliases, preferred_type="date")
        product_col = self.find_column(df, product_aliases)
        revenue_col = self.find_column(df, revenue_aliases, preferred_type="numeric")

        missing = []

        if date_col is None:
            missing.append("date")

        if product_col is None:
            missing.append("product/store/category/customer")

        if revenue_col is None:
            missing.append("revenue/sales amount")

        if missing:
            raise ValueError(
                f"Could not identify required columns: {missing}. "
                f"Available columns: {list(df.columns)}"
            )

        df = df.rename(
            columns={
                date_col: "date",
                product_col: "product",
                revenue_col: "revenue"
            }
        )

        df["date"] = pd.to_datetime(df["date"], errors="coerce")
        df["revenue"] = pd.to_numeric(df["revenue"], errors="coerce")

        df = df.dropna(subset=["date", "product", "revenue"])

        if df.empty:
            raise ValueError(
                "After cleaning, no valid rows remain. "
                "Please check date, product/store/category/customer, and revenue columns."
            )

        return df